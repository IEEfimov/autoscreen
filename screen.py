import pyautogui as gui
import os
import time
from docx import Document


def doMagic(regions):
    """Screen of Array of Regions and
        returns docx.Document"""
    reg1 = regions[0]
    reg2 = regions[1]
    reg3 = regions[2]

    document = Document("resource/template.docx")
    table = document.tables[0]
    pos = 0;
    isChanged = 1

    gui.click(x=40, y=65)
    time.sleep(0.5)
    gui.click(x=93, y=65)
    time.sleep(0.5)

    for i in range(0, 12):
        pic10 = gui.screenshot(region=(reg1[0], reg1[1] + 11 * pos, reg1[2], 11))
        pic0 = pic10
        while 1:
             pic0 = gui.screenshot(region=(reg1[0], reg1[1] + 11 * pos, reg1[2], 11))
             print("{} {}".format(pos, isChanged))
             color = pic0.getpixel((1, 5))
             if color == (192, 192, 192): break

             pos += 1
             isChanged += 1;
             if isChanged > 50:
                 isChanged = 1
                 pos = 0;
        pic1 = gui.screenshot(region=reg2)
        pic2 = gui.screenshot(region=reg3)
        pic10.save('resource/temp0.png')
        pic10 = pic0;
        pic1.save('resource/temp1.png')
        pic2.save('resource/temp2.png')
        row = table.add_row()
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p2 = row.cells[2].paragraphs[0]
        r0 = p0.add_run()
        r0.add_picture('resource/temp0.png')
        r1 = p1.add_run()
        r1.add_picture('resource/temp1.png')
        r2 = p2.add_run()
        r2.add_picture('resource/temp2.png')

        if (pos > 2):
            isChanged = 0;
        gui.click(x=140, y=65)

        time.sleep(0.05)

    name = str(time.strftime('%H%M%S', time.gmtime()))
    document.save('{}.docx'.format(name))
    os.system('{}.docx'.format(name))

    print(pos)

# class Region:
#     left = 0
#     right = 0
#     width = 100
#     height = 100
#     image = None
#
#     def __init__(self, left, right, width, height):
#         self.left = left
#         self.right = right
#         self.width = width
#         self.height = height
#         return
